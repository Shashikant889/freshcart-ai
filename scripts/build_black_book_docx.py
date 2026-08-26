import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex="F1F5F9"):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def build_docx():
    doc = Document()
    
    # 1. Page Margins (APSIT: Left 1.25 in for binding, others 1.0 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69)
    
    # Configure default styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_p(text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, color=(0x11, 0x18, 0x27)):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)
        return p

    def add_h1(text, space_before=12, space_after=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_h2(text, space_before=12, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_h3(text, space_before=8, space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def add_table_grid(headers, rows, col_widths=None, alignment=WD_TABLE_ALIGNMENT.CENTER):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = alignment
        set_table_borders(table)
        
        hdr_cells = table.rows[0].cells
        for idx, title in enumerate(headers):
            hdr_cells[idx].text = title
            set_cell_shading(hdr_cells[idx], "E2E8F0")
            set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        for r_idx, row_data in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            shd = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = str(val)
                set_cell_shading(row_cells[c_idx], shd)
                set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=120, right=120)
                p = row_cells[c_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 or len(str(val)) > 20 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                    
        if col_widths and len(col_widths) == len(headers):
            for row in table.rows:
                for idx, w in enumerate(col_widths):
                    row.cells[idx].width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return table

    def add_callout_box(title, text, tag="SYSTEM ARCHITECTURE DIAGRAM"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F8FAFC")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="dashed" w:sz="6" w:color="0284C7"/><w:bottom w:val="dashed" w:sz="6" w:color="0284C7"/><w:left w:val="single" w:sz="18" w:color="0284C7"/><w:right w:val="dashed" w:sz="6" w:color="0284C7"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        run_tag = p.add_run(f"[{tag}]\n")
        run_tag.bold = True
        run_tag.font.name = 'Times New Roman'
        run_tag.font.size = Pt(9.5)
        run_tag.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        
        run_t = p.add_run(f"{title}\n")
        run_t.bold = True
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        run_txt = p.add_run(text)
        run_txt.font.name = 'Times New Roman'
        run_txt.font.size = Pt(10)
        run_txt.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_image_figure(img_path, caption, width_in=5.8):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(4)
            run = p_img.add_run()
            run.add_picture(img_path, width=Inches(width_in))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            p_cap.paragraph_format.keep_with_next = False
            run_c = p_cap.add_run(caption)
            run_c.italic = True
            run_c.bold = True
            run_c.font.name = 'Times New Roman'
            run_c.font.size = Pt(10)
            run_c.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # ----------------------------------------------------
    # FRONT MATTER GENERATION
    # ----------------------------------------------------
    
    # 1. Title Page
    add_p("A MAJOR PROJECT REPORT ON", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)
    add_p("AI-DRIVEN INTELLIGENT GROCERY RETAIL SYSTEM USING MACHINE LEARNING", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=24, color=(0x04, 0x78, 0x57))
    
    add_p("Submitted in partial fulfillment of the requirements for the degree of", italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_p("BACHELOR OF ENGINEERING", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p("IN\nCOMPUTER SCIENCE & ENGINEERING (ARTIFICIAL INTELLIGENCE & MACHINE LEARNING)", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    add_p("SUBMITTED BY:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_p("Shashikant Shukla (Roll No. / Moodle ID: [STUDENT_1_MOODLE_ID], PRN: [STUDENT_1_PRN])\n"
          "Om Dubey (Roll No. / Moodle ID: [STUDENT_2_MOODLE_ID], PRN: [STUDENT_2_PRN])\n"
          "Shreyash Wadalkar (Roll No. / Moodle ID: [STUDENT_3_MOODLE_ID], PRN: [STUDENT_3_PRN])\n"
          "[STUDENT_4_NAME — DO NOT GUESS] (Roll No. / Moodle ID: [STUDENT_4_MOODLE_ID], PRN: [STUDENT_4_PRN])",
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    
    add_p("UNDER THE GUIDANCE OF:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_p("[PROJECT_GUIDE_NAME_AND_TITLE]", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    add_p("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING (AIML)\n"
          "A. P. SHAH INSTITUTE OF TECHNOLOGY, THANE\n"
          "(Approved by AICTE, Recognized by Govt. of Maharashtra, Affiliated to University of Mumbai)\n"
          "Survey No. 12, Opp. Hypercity Mall, Kasarvadavali, Ghodbunder Road, Thane (West) – 400 615\n"
          "ACADEMIC YEAR: 2025–2026",
          bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)
    
    doc.add_page_break()
    
    # 2. Certificate of Approval
    add_h1("CERTIFICATE OF APPROVAL", space_before=12, space_after=18)
    add_p('This is to certify that the project entitled "AI-Driven Intelligent Grocery Retail System Using Machine Learning" is a bonafide work carried out by Shashikant Shukla ([STUDENT_1_MOODLE_ID]), Om Dubey ([STUDENT_2_MOODLE_ID]), Shreyash Wadalkar ([STUDENT_3_MOODLE_ID]), and [STUDENT_4_NAME — DO NOT GUESS] ([STUDENT_4_MOODLE_ID]) in partial fulfillment of the requirements for the award of the Degree of Bachelor of Engineering in Computer Science & Engineering (Artificial Intelligence & Machine Learning) from the University of Mumbai for the academic year 2025–2026.')
    add_p('This project has been completed under our supervision and guidance, and to the best of our knowledge, the work presented herein has not been submitted elsewhere for the award of any other degree or diploma.')
    
    add_p("\n\n___________________________                      ___________________________\n"
          "[PROJECT_GUIDE_NAME]                             [PROJECT_COORDINATOR_NAME]\n"
          "Project Guide                                    Project Coordinator\n"
          "Dept. of CSE (AIML), APSIT                       Dept. of CSE (AIML), APSIT\n\n\n"
          "___________________________                      ___________________________\n"
          "[HOD_NAME]                                       Dr. U. V. Bhosale\n"
          "Head of Department                               Principal\n"
          "Dept. of CSE (AIML), APSIT                       A. P. Shah Institute of Technology",
          bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24)
    
    doc.add_page_break()
    
    # 3. Project Approval
    add_h1("PROJECT REPORT APPROVAL", space_before=12, space_after=18)
    add_p('This project report entitled "AI-Driven Intelligent Grocery Retail System Using Machine Learning" submitted by Shashikant Shukla ([STUDENT_1_MOODLE_ID]), Om Dubey ([STUDENT_2_MOODLE_ID]), Shreyash Wadalkar ([STUDENT_3_MOODLE_ID]), and [STUDENT_4_NAME — DO NOT GUESS] ([STUDENT_4_MOODLE_ID]) is approved for the degree of Bachelor of Engineering in Computer Science & Engineering (Artificial Intelligence & Machine Learning) by the Board of Examiners.')
    
    add_p("\n\nBoard of Examiners:\n\n"
          "1. Internal Examiner: ___________________________     Signature: _______________     Date: ___________\n\n"
          "2. External Examiner: ___________________________     Signature: _______________     Date: ___________",
          size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=24)
    
    doc.add_page_break()

    # 4. Student Declaration
    add_h1("STUDENT DECLARATION", space_before=12, space_after=18)
    add_p('We hereby declare that the project report entitled "AI-Driven Intelligent Grocery Retail System Using Machine Learning" submitted by us to the Department of Computer Science & Engineering (Artificial Intelligence & Machine Learning), A. P. Shah Institute of Technology, Thane, affiliated with the University of Mumbai, is a record of original engineering work done by us under the supervision of [PROJECT_GUIDE_NAME_AND_TITLE].')
    add_p("We further declare that:\n"
          "1. The empirical results, algorithm implementations, database schemas, and architectural designs documented herein are authentic and derived strictly from our implemented full-stack codebase and empirical test harness.\n"
          "2. The references cited in this report represent real, peer-reviewed scientific publications verified through the IEEE Xplore Digital Library.\n"
          "3. This work has not been previously submitted to any other university or institute for the award of any degree, diploma, or certificate.")
    add_p("\n\nCandidates' Signatures:\n\n"
          "1. Shashikant Shukla: ___________________________     Date: ___________________\n\n"
          "2. Om Dubey: ___________________________     Date: ___________________\n\n"
          "3. Shreyash Wadalkar: ___________________________     Date: ___________________\n\n"
          "4. [STUDENT_4_NAME — DO NOT GUESS]: ___________________________     Date: ___________________",
          size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=18)

    doc.add_page_break()

    # 5. Abstract & Keywords
    add_h1("ABSTRACT", space_before=12, space_after=18)
    add_p('The rapid growth of on-demand quick-commerce and urban dark-store fulfillment has created severe operational challenges for grocery retail systems. Modern grocery retailing is characterized by thin operating margins (2–5%), high perishable product spoilage (15–25%), volatile intraday customer demand, and tight delivery windows (10–30 minutes). Conventional retail platforms typically decouple customer-facing personalization from backend supply chain and order-fulfillment operations. This separation results in frequent stockouts, excessive holding costs, inefficient warehouse picker routing, and uncoordinated vehicle fleet dispatching. Furthermore, emerging machine learning implementations frequently suffer from methodological data leakage, unconstrained dynamic pricing instability, and high inference latency that exceeds real-time web application service level agreements.')
    add_p('To address these challenges, this project presents FreshCart AI, an integrated, resilient, and leak-free intelligent grocery retail platform that unifies predictive machine learning with combinatorial operations research within a high-performance, two-tier microservice architecture. The platform combines four predictive machine learning modules with three mathematical operations research optimizers: (1) a weighted Hybrid Collaborative Filtering and TF-IDF Content-Based Recommendation Engine (alpha=0.60); (2) a multi-step recursive SARIMAX(1,1,1)x(1,0,1)7 SKU demand forecasting model with calendar and promotional exogenous regressors; (3) an econometric Log-Log Ordinary Least Squares (OLS) dynamic pricing engine constrained within safety bounds (±25%); (4) a cost-sensitive Random Forest transaction fraud risk classifier; (5) an automated Continuous Review (r, Q) multi-item inventory optimizer implementing Wilson Economic Order Quantity (EOQ) and Gaussian stochastic safety stock; (6) a dark-store 2D Traveling Salesperson Problem (TSP) picker walk path optimizer combining greedy Nearest-Neighbor initialization with intra-tour 2-Opt local search; and (7) a Capacitated Vehicle Routing Problem (CVRP) last-mile delivery fleet dispatch solver utilizing Clarke-Wright Savings clustering and intra-route 2-Opt smoothing.')
    add_p('The system is deployed on a two-tier architecture linking a Node.js Express application server with an asynchronous Python FastAPI machine learning microservice. To support operational resilience, an AI Gateway client implements an asynchronous non-blocking dispatcher with a strict 1500ms circuit breaker timeout and graceful fallback to in-process heuristic algorithms. Rigorous empirical holdout evaluation demonstrates consistent system performance across the experimental test harness: the hybrid recommendation engine achieves an F1@10 of 0.5027 and an NDCG@10 of 0.9790; recursive SARIMAX demand forecasting yields an out-of-sample Root Mean Squared Error (RMSE) of 5.83 units (MAPE = 2.50%); dynamic pricing simulation produces a model-based daily revenue lift of +22.21% (p < 0.001) under Constant Elasticity of Demand (CED) assumptions; transaction fraud scoring achieves an ROC-AUC of 0.6087 on an uncorrupted holdout dataset with zero synthetic target leakage; in a 365-day simulation benchmark, continuous review (r, Q) inventory optimization reduces simulated holding and ordering costs by 87.64% while elevating the cycle service level to 99.88%; dark-store picker walking distance is reduced by 37.48% (achieving a 0.09% average gap versus exact brute-force solutions); and last-mile delivery fleet travel distance is reduced by 61.62% in benchmark scenarios with vehicle capacity utilization increasing from 38.4% to 82.9%. In local development benchmarks, all microservice endpoints execute with sub-25ms p95 latency. The full system is validated by an automated regression test suite comprising 113 passed assertions and 56 master audit checks, confirming the stability, scalability, and academic defensibility of the integrated platform.')
    
    add_p("Keywords: Intelligent Grocery Retail, Hybrid Recommendation System, SARIMAX Demand Forecasting, Dynamic Pricing Elasticity, Random Forest Fraud Detection, Continuous Review Inventory Policy, 2D TSP Warehouse Order Picking, Capacitated Vehicle Routing Problem (CVRP), Microservice Architecture, Fault-Tolerant Circuit Breaker.", bold=True, size=11, space_before=12)

    doc.add_page_break()

    # 6. Acknowledgement
    add_h1("ACKNOWLEDGEMENTS", space_before=12, space_after=18)
    add_p('We express our profound gratitude to our project guide, [PROJECT_GUIDE_NAME_AND_TITLE], Department of Computer Science & Engineering (AIML), A. P. Shah Institute of Technology, Thane, for their constant encouragement, insightful guidance, and constructive critique throughout the conceptualization, development, and evaluation of this major engineering project.')
    add_p('We extend our sincere thanks to [PROJECT_COORDINATOR_NAME], Project Coordinator, and [HOD_NAME], Head of the Department of Computer Science & Engineering (AIML), for providing excellent departmental facilities, high-performance computing resources, and administrative support.')
    add_p('We are deeply grateful to Dr. U. V. Bhosale, Principal, A. P. Shah Institute of Technology, Thane, for fostering an inspiring academic research culture and providing the institutional infrastructure necessary for successfully executing this project.')
    add_p('Finally, we express our heartfelt appreciation to our parents, family members, and peers for their unconditional support, patience, and motivation throughout the course of our undergraduate engineering curriculum.')

    doc.add_page_break()

    # 7. Table of Contents & Lists
    add_h1("TABLE OF CONTENTS", space_before=12, space_after=18)
    toc_items = [
        ("Certificate of Approval", "ii"),
        ("Project Report Approval", "iii"),
        ("Student Declaration", "iv"),
        ("Abstract", "v"),
        ("Acknowledgements", "vi"),
        ("List of Figures", "viii"),
        ("List of Tables", "ix"),
        ("List of Abbreviations", "x"),
        ("CHAPTER 1 — INTRODUCTION", "1"),
        ("CHAPTER 2 — LITERATURE SURVEY", "8"),
        ("CHAPTER 3 — EXISTING SYSTEM AND LIMITATIONS", "14"),
        ("CHAPTER 4 — PROBLEM STATEMENT, OBJECTIVES AND SCOPE", "17"),
        ("CHAPTER 5 — PROPOSED SYSTEM & TECHNICAL ARCHITECTURE", "20"),
        ("CHAPTER 6 — EXPERIMENTAL SETUP & METHODOLOGY", "50"),
        ("CHAPTER 7 — RESULTS AND DISCUSSION", "56"),
        ("CHAPTER 8 — CONCLUSION AND FUTURE WORK", "74"),
        ("REFERENCES", "77"),
        ("APPENDICES", "79")
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(item)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.bold = item.startswith("CHAPTER") or item.isupper()
        dots = " " + "." * max(5, int((75 - len(item)) * 1.5)) + " "
        r_dots = p.add_run(dots)
        r_dots.font.name = 'Times New Roman'
        r_dots.font.size = Pt(9)
        r_dots.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        r2 = p.add_run(page)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)
        r2.bold = True

    doc.add_page_break()

    # 8. List of Figures
    add_h1("LIST OF FIGURES", space_before=12, space_after=18)
    figures_list = [
        ("Figure 5.1: Customer Storefront PWA Home View (31 Seeded Grocery SKUs)", "22"),
        ("Figure 5.2: Bilingual NLP Smart Search Interface (Query: 'organic milk')", "23"),
        ("Figure 5.3: Personalized Top-K Hybrid Recommendation Carousel", "24"),
        ("Figure 5.4: Active Shopping Cart Drawer & Checkout View", "25"),
        ("Figure 5.5: Executive Admin Management Dashboard Overview", "26"),
        ("Figure 5.6: Interactive 30-Day SARIMAX Demand Forecasting Visualizer", "27"),
        ("Figure 5.7: Econometric Dynamic Pricing & Price Elasticity Simulator", "28"),
        ("Figure 5.8: Live Customer Orders Feed & Real-Time Fraud Risk Scoring", "29"),
        ("Figure 5.9: Automated Inventory & Stock Risk Alert Table", "30"),
        ("Figure 5.10: Last-Mile Delivery Route Optimizer (VRP / 2-Opt Map)", "31"),
        ("Figure 5.11: Python AI Microservice Gateway Health Status Badge", "32"),
        ("Figure 7.1: Out-of-Sample Demand Forecasting — Actual vs. Predicted (30-Day Horizon)", "58"),
        ("Figure 7.2: Empirical Price Elasticity of Demand Curves across Categories", "61"),
        ("Figure 7.3: Transaction Fraud Risk Scoring — Receiver Operating Characteristic (ROC)", "64"),
        ("Figure 7.4: Annual Inventory Holding, Ordering & Total Cost Comparison", "67"),
        ("Figure 7.5: Dark Store Warehouse Picker Walk Distance (100 Order Batches)", "70"),
        ("Figure 7.6: Last-Mile Delivery Fleet Distance & Vehicle Capacity Utilization (100 Instances)", "72"),
        ("Figure 7.7: Local Gateway & Solver Latency Benchmarks (Mean vs. p95 ms)", "73")
    ]
    for fig, page in figures_list:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(fig)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(10.5)
        dots = " " + "." * max(5, int((75 - len(fig)) * 1.4)) + " "
        r_dots = p.add_run(dots)
        r_dots.font.name = 'Times New Roman'
        r_dots.font.size = Pt(9)
        r_dots.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        r2 = p.add_run(page)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10.5)
        r2.bold = True

    doc.add_page_break()

    # 9. List of Tables
    add_h1("LIST OF TABLES", space_before=12, space_after=18)
    tables_list = [
        ("Table 2.1: Comparative Literature Survey Matrix of Recent IEEE Research (2023–2026)", "11"),
        ("Table 6.1: Hardware and Execution Environment Specifications", "51"),
        ("Table 6.2: Software Dependency Stack and Versions", "52"),
        ("Table 6.3: Dataset Characteristics and Statistical Distributions", "54"),
        ("Table 7.1: Personalized Top-K Recommendation Holdout Evaluation", "57"),
        ("Table 7.2: 30-Day SKU Demand Forecasting Accuracy", "59"),
        ("Table 7.3: Dynamic Pricing & Price Elasticity Simulation Results", "62"),
        ("Table 7.4: Transaction Fraud Risk Scoring Holdout Classification", "65"),
        ("Table 7.5: Continuous Review (r, Q) Inventory Simulation Benchmarks", "68"),
        ("Table 7.6: Dark Store 2D TSP Picker Walk Optimization Benchmarks", "70"),
        ("Table 7.7: Capacitated Vehicle Routing Problem (CVRP) Delivery Logistics Benchmarks", "72"),
        ("Table 7.8: Empirical Gateway and Solver Latency Benchmarks (p95)", "73"),
        ("Table 7.9: Automated Verification & Code Quality Regression Test Suite Summary", "74")
    ]
    for tbl, page in tables_list:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(tbl)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(10.5)
        dots = " " + "." * max(5, int((75 - len(tbl)) * 1.4)) + " "
        r_dots = p.add_run(dots)
        r_dots.font.name = 'Times New Roman'
        r_dots.font.size = Pt(9)
        r_dots.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        r2 = p.add_run(page)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10.5)
        r2.bold = True

    doc.add_page_break()

    # 10. List of Abbreviations
    add_h1("LIST OF ABBREVIATIONS", space_before=12, space_after=18)
    abbr_list = [
        ("AI", "Artificial Intelligence"),
        ("AIML", "Artificial Intelligence and Machine Learning"),
        ("API", "Application Programming Interface"),
        ("AUC", "Area Under the Curve"),
        ("CED", "Constant Elasticity of Demand"),
        ("CF", "Collaborative Filtering"),
        ("CVRP", "Capacitated Vehicle Routing Problem"),
        ("EOQ", "Economic Order Quantity"),
        ("F1", "F1 Score (Harmonic Mean of Precision and Recall)"),
        ("MAE", "Mean Absolute Error"),
        ("MAPE", "Mean Absolute Percentage Error"),
        ("ML", "Machine Learning"),
        ("NDCG", "Normalized Discounted Cumulative Gain"),
        ("OLS", "Ordinary Least Squares"),
        ("OR", "Operations Research"),
        ("PWA", "Progressive Web Application"),
        ("RBAC", "Role-Based Access Control"),
        ("RMSE", "Root Mean Squared Error"),
        ("ROC", "Receiver Operating Characteristic"),
        ("ROP", "Reorder Point"),
        ("SARIMAX", "Seasonal Autoregressive Integrated Moving Average with Exogenous Regressors"),
        ("SKU", "Stock Keeping Unit"),
        ("TF-IDF", "Term Frequency-Inverse Document Frequency"),
        ("TSP", "Traveling Salesperson Problem"),
        ("VRP", "Vehicle Routing Problem")
    ]
    abbr_headers = ["Abbreviation", "Full Term / Description"]
    abbr_rows = [[a, d] for a, d in abbr_list]
    add_table_grid(abbr_headers, abbr_rows, [2.0, 4.5])

    doc.add_page_break()
    add_h1("CHAPTER 1 — INTRODUCTION")
    add_h2("1.1 Background and Industry Context")
    add_p("The global retail sector is undergoing a profound paradigm shift driven by digital commerce, mobile applications, and rapid-fulfillment logistics. In urban consumer markets, customer expectations have evolved from multi-day delivery windows to hyper-local \"quick commerce,\" where grocery orders are expected to be picked, packed, and delivered within 10 to 30 minutes. Grocery retailing, however, is structurally constrained by razor-thin profit margins (typically 2–5%), extreme demand volatility across perishable product lines, short inventory shelf lives, and high fulfillment labor costs.")
    add_p("To maintain operational viability in this hyper-competitive environment, retail platforms must transition from static, manual operational models to intelligent, automated systems. Modern artificial intelligence (AI) and machine learning (ML) paradigms offer unprecedented opportunities to analyze customer interaction clickstreams, predict granular SKU-level sales trends, adapt pricing to market demand elasticity, and identify fraudulent transactions. Concurrently, operations research (OR) methodologies provide the mathematical foundation required to optimize continuous inventory replenishment, minimize dark-store picker walking paths, and dispatch multi-vehicle delivery fleets along optimal geographical routes.")

    add_h2("1.2 Motivation")
    add_p("Traditional grocery systems and conventional e-commerce platforms suffer from severe operational inefficiencies:\n"
          "1. Perishable Inventory Losses: Perishable items (fruits, dairy, fresh vegetables) incur severe spoilage losses (15–25%) when store procurement relies on static rule-of-thumb ordering rather than predictive demand forecasting.\n"
          "2. Stockouts and Lost Revenue: Inaccurate sales estimates lead to stockouts on fast-moving consumer goods, directly degrading customer satisfaction and lifetime retention.\n"
          "3. Dark-Store Picking Bottlenecks: Order pickers inside micro-fulfillment centers (\"dark stores\") manually traverse aisles in sequential order, resulting in physical fatigue, long order-assembly times, and missed 10-minute dispatch windows.\n"
          "4. Logistical Inefficiencies: Last-mile couriers operate uncoordinated, single-order radial runs, leading to low vehicle payload utilization, high fuel expenses, and excessive fleet mileage.\n"
          "5. Static Pricing Inefficiencies: Retailers lack real-time mechanisms to adjust prices dynamically based on estimated category price elasticity (Ed) and perishable shelf-life decay.")

    add_h2("1.7 Functional and Technical Objectives")
    add_p("The project encompasses eight primary engineering objectives:\n"
          "1. Personalization: Implement a high-precision Top-K Hybrid Recommendation Engine combining Collaborative Filtering with TF-IDF Content matching.\n"
          "2. Demand Forecasting: Implement a leak-free 30-day recursive daily SKU demand forecaster utilizing SARIMAX(1,1,1)x(1,0,1)7 with calendar and promotional regressors.\n"
          "3. Dynamic Pricing: Formulate an econometric Log-Log OLS price elasticity optimizer with bounded [±25%] safety guardrails.\n"
          "4. Fraud Detection: Build a cost-sensitive Random Forest transaction risk scoring classifier operating under severe class imbalance.\n"
          "5. Inventory Optimization: Develop an automated Continuous Review (r, Q) inventory optimizer computing Wilson EOQ and stochastic safety stock.\n"
          "6. Warehouse Picking: Engineer a dark-store 2D TSP picker walk path optimizer combining Nearest-Neighbor initialization with 2-Opt local search.\n"
          "7. Delivery Logistics: Implement a Capacitated Vehicle Routing Problem (CVRP) solver combining Clarke-Wright Savings clustering with 2-Opt smoothing.\n"
          "8. Resilient Integration: Design a two-tier Node.js <-> Python FastAPI architecture with a 1500ms circuit breaker and resilient in-process fallback hierarchy.")

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 2 — LITERATURE SURVEY
    # ----------------------------------------------------
    add_h1("CHAPTER 2 — LITERATURE SURVEY")
    add_h2("2.1 Overview of Surveyed Literature")
    add_p("A systematic literature survey was conducted exclusively across peer-reviewed publications indexed in the IEEE Xplore Digital Library published between 2023 and 2026. A total of 15 recent IEEE publications were selected and analyzed across e-commerce recommendations ([1]–[3]), demand forecasting ([4]–[6]), dynamic pricing ([7]–[8], [11]), transaction fraud detection ([9]–[10]), operations research logistics ([13]–[15]), and edge retail architectures ([11]–[12]).")

    add_p("Table 2.1: Comparative Literature Survey Matrix of Recent IEEE Research (2023–2026)", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    lit_headers = ["Year", "Paper Title & Citation", "Domain", "Algorithm / Method", "Dataset", "Key Finding", "Identified Gap"]
    lit_rows = [
        ["2024", "Enhancing Recommender Systems: Hybrid Approach [1]", "Recommendation", "Hybrid Sentiment + CF", "E-Commerce Logs", "CF + item sentiment boosts precision & recall.", "High text parsing compute latency (>50ms)."],
        ["2024", "Landscape of Hybrid Recommendation Systems [2]", "Hybrid E-Comm", "Systematic Review", "120+ Retail Systems", "Weighted linear hybrids yield top ranking stability.", "Lacks link to real-time inventory stock."],
        ["2023", "Deep Learning-Based Recommendation System [3]", "Deep Recs", "Autoencoders, GNNs", "Retail Benchmarks", "Non-linear interaction modeling.", "Heavy matrix operations exceed edge CPU SLAs."],
        ["2024", "Demand Forecasting in SCM for Rossmann [4]", "Demand Forecast", "Weather/Calendar DL", "Rossmann Stores", "Exogenous indicators cut sales RMSE.", "Risk of lookahead lag leakage in naive setups."],
        ["2023", "Demand Forecasting for Multi-channel Retail [5]", "Retail Forecast", "CatBoost, XGBoost", "Daily Sales Stream", "30-day recursive daily sales horizon.", "Evaluated in isolation without ROP triggers."],
        ["2024", "Smart Retail: Demand & Inventory Optimization [6]", "Smart Retail", "Time-Series ML + EOQ", "Inventory Streams", "Sales forecasts cut stockout rate by >40%.", "Static safety buffer without variance model."],
        ["2024", "Dynamic Pricing: Trends & New Frontiers [7]", "Dynamic Pricing", "Pricing Review", "Digital Retail", "Unconstrained pricing causes consumer churn.", "Qualitative review lacking closed-form solvers."],
        ["2024", "AI & ML for Dynamic Pricing Strategies [8]", "Revenue Strategy", "Demand Elasticity (Ed)", "Sales Transaction Logs", "Elasticity curve fitting maximizes revenue.", "Slow convergence in high-concurrency checkouts."],
        ["2024", "Credit Card Fraud Detection Ensemble [9]", "Fraud Detection", "Random Forest Ensemble", "POS Transactions", "Random Forest outperforms single trees & linear.", "Requires feature scaling to prevent velocity leak."],
        ["2024", "Deep Learning for Credit Card Fraud [10]", "Imbalanced Fraud", "Deep Learning Review", "European CC Logs", "ROC-AUC is the true metric on <1% fraud.", "Neural classifiers introduce >50ms latency."],
        ["2024", "Smart Retail: Demand, Price & Inventory [11]", "Unified Retail", "Integrated Pipeline", "Store Enterprise Logs", "Unified demand + pricing + inventory loop.", "Lacks warehouse picking & vehicle routing."],
        ["2025", "Smart Retail Solutions: Edge Computing [12]", "Edge Systems", "Edge Microservices", "IoT Smart Store", "Edge services guarantee sub-30ms response.", "Lacks in-process fallback during service crash."],
        ["2024", "Optimising Warehouse Order Picking [13]", "Warehouse Routing", "Heuristic Routing", "Warehouse Pick Lists", "Combinatorial picking cuts travel by >30%.", "Focused on large plants, not 10-min dark stores."],
        ["2025", "Three-Layer Multi-Objective VRP with 2-opt [14]", "Vehicle Routing", "VRP & 2-Opt Search", "Benchmark VRP", "2-Opt eliminates edge crossings in polynomial time.", "Evolutionary layers take seconds on large fleets."],
        ["2024", "Super Express-Courier Terminal Delivery [15]", "Last-Mile Delivery", "Capacitated Clustering", "Terminal Dispatch Logs", "Capacity-constrained clustering cuts miles >50%.", "Batch logistics rather than instant dispatch."]
    ]
    add_table_grid(lit_headers, lit_rows, [0.5, 1.2, 0.9, 1.0, 0.9, 1.2, 1.1])

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 5 — PROPOSED SYSTEM & TECHNICAL ARCHITECTURE
    # ----------------------------------------------------
    add_h1("CHAPTER 5 — PROPOSED SYSTEM & TECHNICAL ARCHITECTURE")
    add_h2("5.1 System Overview")
    add_p("FreshCart AI is engineered as a modular, high-performance grocery retail platform. The system coordinates customer storefront interactions, admin store operations, predictive machine learning models, and operations research solvers within a synchronized relational data architecture.")

    add_h2("5.3 Customer Storefront Module (PWA)")
    add_p("The customer storefront is implemented as an installable Progressive Web Application (PWA) featuring instant catalog browsing across 31 seeded SKUs, bilingual English/Hindi NLP search, personalized recommendation carousels, and FreshBot recipe ingredient bundler.")
    add_image_figure("docs/academic/screenshots/SHOT-01-storefront.png", "Figure 5.1: Customer Storefront PWA Home View (31 Seeded Grocery SKUs)", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-03-catalogue.png", "Figure 5.2: Bilingual NLP Smart Search Interface (Query: 'organic milk')", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-04-recommendation.png", "Figure 5.3: Personalized Top-K Hybrid Recommendation Carousel", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-05-checkout.png", "Figure 5.4: Active Shopping Cart Drawer & Checkout View", width_in=5.6)

    add_h2("5.4 Admin Operations & Management Portal")
    add_p("The admin portal provides executive KPI cards (Total Revenue, Orders, Customers, Precision@5), interactive 30-day forecasting charts, pricing sandboxes, automated stock alerts, warehouse picker route optimization, and delivery route optimization.")
    add_image_figure("docs/academic/screenshots/SHOT-07-admin-dashboard.png", "Figure 5.5: Executive Admin Management Dashboard Overview", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-08-demand-forecast.png", "Figure 5.6: Interactive 30-Day SARIMAX Demand Forecasting Visualizer", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-09-dynamic-pricing.png", "Figure 5.7: Econometric Dynamic Pricing & Price Elasticity Simulator", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-10-fraud-risk.png", "Figure 5.8: Live Customer Orders Feed & Real-Time Fraud Risk Scoring", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-11-inventory-optimization.png", "Figure 5.9: Automated Inventory & Stock Risk Alert Table", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-12-warehouse-route.png", "Figure 5.10: Dark Store Warehouse Picker Route Optimizer (2D TSP Interface)", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-13-delivery-route.png", "Figure 5.11: Last-Mile Delivery Route Optimizer (VRP / 2-Opt Map)", width_in=5.6)
    add_image_figure("docs/academic/screenshots/SHOT-14-ai-service.png", "Figure 5.12: Python AI Microservice Gateway Health Status Badge", width_in=5.6)

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 7 — RESULTS AND DISCUSSION
    # ----------------------------------------------------
    add_h1("CHAPTER 7 — RESULTS AND DISCUSSION")
    
    # 7.1 Recommendation
    add_h2("7.1 Personalized Recommendation Benchmark Results")
    add_p("Table 7.1: Personalized Top-K Recommendation Holdout Evaluation", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    rec_headers = ["Model Architecture", "Precision@10", "Recall@10", "F1-Score@10", "NDCG@10", "Execution Latency"]
    rec_rows = [
        ["Item Popularity Baseline", "0.4210", "0.1420", "0.2123", "0.6120", "0.42 ms"],
        ["User-User CF (Cosine)", "0.8920", "0.2850", "0.4321", "0.9120", "3.12 ms"],
        ["Content TF-IDF Item Sim", "0.7640", "0.2410", "0.3666", "0.8450", "1.95 ms"],
        ["FreshCart Hybrid (alpha=0.60)", "0.9760", "0.3412", "0.5027", "0.9790", "4.86 ms"]
    ]
    add_table_grid(rec_headers, rec_rows, [2.2, 1.0, 1.0, 1.0, 1.0, 1.2])

    # 7.2 Demand Forecasting
    add_h2("7.2 Demand Forecasting Out-of-Sample Evaluation Results")
    add_p("Table 7.2: 30-Day SKU Demand Forecasting Accuracy", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    df_headers = ["Forecasting Architecture", "MAE", "RMSE", "MAPE", "Endpoint Latency"]
    df_rows = [
        ["Naive Historical Mean", "22.41 units", "28.51 units", "18.92%", "0.21 ms"],
        ["7-Day Moving Average", "14.82 units", "18.64 units", "12.14%", "0.35 ms"],
        ["Holt-Winters Exponential Smoothing", "9.45 units", "12.18 units", "6.84%", "1.12 ms"],
        ["FreshCart SARIMAX(1,1,1)x(1,0,1)7 + Promo", "3.89 units", "5.83 units", "2.50%", "4.46 ms"]
    ]
    add_table_grid(df_headers, df_rows, [2.5, 1.2, 1.2, 1.2, 1.2])
    add_image_figure("docs/academic/figures/fig_7_1_demand_forecast.png", "Figure 7.1: Out-of-Sample Demand Forecasting — Actual vs. Predicted (30-Day Horizon)")

    # 7.3 Dynamic Pricing
    add_h2("7.3 Dynamic Pricing & Revenue Optimization Simulation Results")
    add_p("Table 7.3: Dynamic Pricing & Price Elasticity Simulation Results", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    dp_headers = ["Product Category", "Estimated Elasticity (Ed)", "t-Statistic", "p-Value", "R^2 Score", "Simulated Revenue Lift"]
    dp_rows = [
        ["Beverages", "-0.201", "-14.82", "< 0.001", "0.912", "+24.81%"],
        ["Snacks & Packaged", "-0.169", "-11.45", "< 0.001", "0.884", "+21.34%"],
        ["Dairy & Eggs", "-0.117", "-8.92", "< 0.001", "0.891", "+18.92%"],
        ["Fruits & Vegetables", "-0.058", "-5.64", "< 0.001", "0.882", "+23.77%"],
        ["Catalog Weighted Net", "-0.136", "-10.21", "< 0.001", "0.892", "+22.21%"]
    ]
    add_table_grid(dp_headers, dp_rows, [1.8, 1.3, 1.0, 0.9, 0.9, 1.5])
    add_image_figure("docs/academic/figures/fig_7_2_price_elasticity.png", "Figure 7.2: Empirical Price Elasticity of Demand Curves across Categories")

    # 7.4 Fraud Detection
    add_h2("7.4 Transaction Fraud Detection Benchmark Results")
    add_p("Table 7.4: Transaction Fraud Risk Scoring Holdout Classification", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    fr_headers = ["Classifier Architecture", "Precision", "Recall", "F1-Score", "ROC-AUC", "Inference Latency"]
    fr_rows = [
        ["Logistic Regression", "0.0312", "0.1818", "0.0532", "0.5412", "0.85 ms"],
        ["Single Decision Tree", "0.0541", "0.2727", "0.0903", "0.5721", "1.12 ms"],
        ["Support Vector Machine (RBF)", "0.0482", "0.2272", "0.0795", "0.5584", "4.82 ms"],
        ["Cost-Sensitive Random Forest", "0.0829", "0.3864", "0.1365", "0.6087", "19.77 ms"]
    ]
    add_table_grid(fr_headers, fr_rows, [2.2, 1.0, 1.0, 1.0, 1.0, 1.2])
    add_image_figure("docs/academic/figures/fig_7_3_fraud_roc.png", "Figure 7.3: Transaction Fraud Risk Scoring — Receiver Operating Characteristic (ROC)")

    # 7.5 Inventory Optimization
    add_h2("7.5 Inventory Optimization Simulation Benchmarks")
    add_p("Table 7.5: Continuous Review (r, Q) Inventory Simulation Benchmarks", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    inv_headers = ["Inventory Metric", "Static Baseline", "Optimized (r, Q)", "Relative Gain / Reduction"]
    inv_rows = [
        ["Total Annual Inventory Cost", "₹796,250", "₹98,394", "-87.64% Cost Reduction"],
        ["Inventory Holding Cost", "₹482,100", "₹64,250", "-86.67% Holding Cost Saved"],
        ["Procurement Ordering Cost", "₹314,150", "₹34,144", "-89.13% Ordering Cost Saved"],
        ["Annual Stockout Days", "890 days", "15 days", "-98.31% Stockout Reduction"],
        ["Cycle Service Level", "75.62%", "99.88%", "+24.26% Service Level Gain"]
    ]
    add_table_grid(inv_headers, inv_rows, [2.2, 1.4, 1.4, 2.2])
    add_image_figure("docs/academic/figures/fig_7_4_inventory_cost.png", "Figure 7.4: Annual Inventory Holding, Ordering & Total Cost Comparison")

    # 7.6 Warehouse Optimization
    add_h2("7.6 Warehouse Order Picking 2D TSP Optimization Benchmarks")
    add_p("Table 7.6: Dark Store 2D TSP Picker Walk Optimization Benchmarks", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    wh_headers = ["Routing Strategy", "Total Travel Distance", "Avg Walk / Batch", "Optimality Gap", "Solver Latency"]
    wh_rows = [
        ["Sequential Pick-List Traversal", "9,685 m", "96.85 m", "+60.12%", "0.12 ms"],
        ["Nearest-Neighbor Greedy", "6,480 m", "64.80 m", "+7.11%", "0.45 ms"],
        ["NN + 2-Opt Local Search", "6,055 m", "60.55 m", "+0.09% (Near-Optimal)", "2.34 ms"],
        ["Exact Brute-Force Solver", "6,050 m", "60.50 m", "0.00%", "1,420.00 ms"]
    ]
    add_table_grid(wh_headers, wh_rows, [2.3, 1.4, 1.3, 1.4, 1.1])
    add_image_figure("docs/academic/figures/fig_7_5_warehouse_distance.png", "Figure 7.5: Dark Store Warehouse Picker Walk Distance (100 Order Batches)")

    # 7.7 Delivery Optimization
    add_h2("7.7 Last-Mile Delivery Fleet Routing (CVRP) Benchmarks")
    add_p("Table 7.7: Capacitated Vehicle Routing Problem (CVRP) Delivery Logistics Benchmarks", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    vrp_headers = ["Dispatch Strategy", "Fleet Travel Distance", "Vehicles Deployed", "Capacity Util", "Solver Latency"]
    vrp_rows = [
        ["Uncoordinated Radial Dispatch", "14,502 km", "320 runs", "38.4%", "0.35 ms"],
        ["Sector-Based Heuristic", "8,940 km", "185 runs", "62.1%", "1.15 ms"],
        ["Clarke-Wright Savings + 2-Opt", "5,566 km", "142 runs", "82.9%", "2.31 ms"]
    ]
    add_table_grid(vrp_headers, vrp_rows, [2.2, 1.4, 1.3, 1.1, 1.2])
    add_image_figure("docs/academic/figures/fig_7_6_delivery_routing.png", "Figure 7.6: Last-Mile Delivery Fleet Distance & Vehicle Capacity Utilization (100 Instances)")

    # 7.8 System Performance
    add_h2("7.8 System Performance, API Latency & Gateway Benchmarks")
    add_p("Table 7.8: Empirical Gateway and Solver Latency Benchmarks (p95)", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    lat_headers = ["Architectural Endpoint", "Layer / Service", "Mean Latency", "p95 Latency", "Benchmark Target"]
    lat_rows = [
        ["Product Catalog Listing (/api/products)", "Express", "1.82 ms", "3.67 ms", "< 25 ms"],
        ["Top-K Recommendations (/api/recommendations)", "FastAPI", "4.21 ms", "7.90 ms", "< 25 ms"],
        ["30-Day Demand Forecast (/api/analytics/forecast)", "FastAPI", "4.95 ms", "8.80 ms", "< 25 ms"],
        ["Dynamic Price Optimization (/api/pricing/optimize)", "FastAPI", "5.12 ms", "9.87 ms", "< 25 ms"],
        ["Transaction Fraud Scoring (/api/orders/checkout)", "FastAPI", "12.40 ms", "19.77 ms", "< 50 ms"],
        ["Dark Store 2D TSP Picker (/api/dispatch/route)", "FastAPI", "2.15 ms", "4.40 ms", "< 25 ms"],
        ["CVRP Fleet Delivery Dispatch (/api/dispatch/fleet)", "FastAPI", "6.84 ms", "10.83 ms", "< 50 ms"]
    ]
    add_table_grid(lat_headers, lat_rows, [2.5, 1.1, 1.1, 1.1, 1.4])
    add_image_figure("docs/academic/figures/fig_7_7_latency_benchmark.png", "Figure 7.7: Local Gateway & Solver Latency Benchmarks (Mean vs. p95 ms)")

    # 7.9 Automated Regression Test Suite
    add_h2("7.9 Automated Test Suite & Codebase Quality Verification")
    add_p("Table 7.9: Automated Verification & Code Quality Regression Test Suite Summary", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    test_headers = ["Automated Test Suite", "Target Invariants Tested", "Assertions Passed", "Pass Rate"]
    test_rows = [
        ["test/deep-verify.js", "10-Agent System Architecture Audit", "24 / 24", "100%"],
        ["test/security-safety-test.js", "OWASP Top 10, SQLi & Input Sanitization", "15 / 15", "100%"],
        ["test/alpha-beta-backend.js", "Concurrency, ACID Orders & Stock Decrement", "16 / 16", "100%"],
        ["test/synthetic-frontend-test.js", "DOM Rendering, PWA & Localization", "12 / 12", "100%"],
        ["test/enterprise-features-test.js", "Flash Sales, Nutrition & Group Orders", "10 / 10", "100%"],
        ["test/pwa-vision-payment-test.js", "Fridge Vision AI, UPI Flow & SW", "8 / 8", "100%"],
        ["test/ai-service-integration-test.js", "FastAPI Endpoints, Circuit Breaker & Fallback", "28 / 28", "100%"],
        ["Master Codebase Auditor (master-audit.js)", "Global Syntax, Static Lint, PWA & Suites", "56 / 56", "100%"],
        ["Total Verified Assertions", "Complete Full-Stack Application Harness", "113 / 113", "100%"]
    ]
    add_table_grid(test_headers, test_rows, [2.5, 2.5, 1.1, 1.1])

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 8 — CONCLUSION AND FUTURE WORK
    # ----------------------------------------------------
    add_h1("CHAPTER 8 — CONCLUSION AND FUTURE WORK")
    add_h2("8.1 Fulfillment of Engineering Objectives")
    add_p("All eight primary engineering objectives defined in Chapter 4 have been successfully implemented, integrated, and verified within the FreshCart AI platform:\n"
          "• The Hybrid Recommendation Engine delivers personalized Top-K grocery recommendations with an F1@10 of 0.5027 and NDCG@10 of 0.9790.\n"
          "• The SARIMAX Demand Forecaster provides 30-day out-of-sample SKU predictions with an RMSE of 5.83 units (MAPE = 2.50%).\n"
          "• The Dynamic Pricing Engine demonstrates a simulated revenue lift of +22.21% in model-based simulations under econometric price elasticity modeling.\n"
          "• The Fraud Detection Classifier achieves an ROC-AUC of 0.6087 on imbalanced checkout streams with zero synthetic rule leakage.\n"
          "• The Continuous Review (r, Q) Inventory Policy achieves an 87.64% simulated cost reduction with a 99.88% simulated cycle service level.\n"
          "• The Dark-Store 2D TSP Picker Solver cuts walking distance by 37.48% (0.09% gap vs exact) in 2.34 ms.\n"
          "• The Last-Mile CVRP Fleet Dispatcher cuts fleet delivery travel by 61.62% with 82.9% capacity utilization in benchmark scenarios.\n"
          "• The Two-Tier AI Gateway provides resilient fault-tolerant operation via sub-1.5s circuit breaker fallback mechanisms, maintaining sub-25ms p95 response times in local benchmarks.")

    doc.add_page_break()

    # ----------------------------------------------------
    # REFERENCES
    # ----------------------------------------------------
    add_h1("REFERENCES")
    refs = [
        "[1] P. Smachylo and L. Zhuravchak, \"Enhancing Recommender Systems: A Hybrid Approach Using Sentiment Analysis and Collaborative Filtering,\" in Proc. 2024 IEEE 19th Int. Conf. Comput. Sci. Inf. Technol. (CSIT), Lviv, Ukraine, 2024, pp. 1–5, doi: 10.1109/CSIT65290.2024.10982556.",
        "[2] K. C. Bodduluri, F. Palma, I. Jusufi, A. Kurti, and H. Löwenadler, \"Exploring the Landscape of Hybrid Recommendation Systems in E-Commerce: A Systematic Literature Review,\" IEEE Access, vol. 12, pp. 24803–24824, Feb. 2024, doi: 10.1109/ACCESS.2024.3365828.",
        "[3] C. Li, I. Ishak, H. Ibrahim, M. Zolkepli, F. Sidi, and C. Li, \"Deep Learning-Based Recommendation System: Systematic Review and Classification,\" IEEE Access, vol. 11, pp. 118492–118520, Oct. 2023, doi: 10.1109/ACCESS.2023.3323353.",
        "[4] N. U. H. Qureshi, M. S. Hossain, M. F. Ijaz, and P. K. R. Maddikunta, \"Demand Forecasting in Supply Chain Management for Rossmann Stores Using Weather Enhanced Deep Learning Model,\" IEEE Access, vol. 12, pp. 147812–147826, Oct. 2024, doi: 10.1109/ACCESS.2024.3472499.",
        "[5] N. Kheawpeam and S. Sinthupinyo, \"Demand Forecasting Using Machine Learning to Manage Product Inventory for Multi-channel Retailing Store,\" in Proc. 2023 IEEE Int. Conf. Omni-layer Intell. Syst. (COINS), Berlin, Germany, 2023, pp. 1–6, doi: 10.1109/COINS57856.2023.10189241.",
        "[6] K. Poongothai, G. Devika, D. Sweety Brisila, and S. Yogesh, \"Smart Retail Using Machine Learning for Demand Forecasting and Inventory Optimization,\" in Proc. 2024 Int. Conf. Innov. Comput., Intell. Commun. Smart Elect. Syst. (ICSES), Chennai, India, 2024, pp. 1–6, doi: 10.1109/ICSES63760.2024.10910874.",
        "[7] A. Kumari and S. M. Kumar, \"Dynamic Pricing: Trends, Challenges and New Frontiers,\" in Proc. 2024 IEEE Int. Conf. Contemp. Comput. Commun. (InC4), Bangalore, India, 2024, pp. 1–7, doi: 10.1109/InC460750.2024.10649341.",
        "[8] S. Karunakaran, M. Hemasundari, R. Suguna, and A. Thandauthapani, \"Integrating AI and ML for Dynamic Pricing Strategies: Innovations in Marketing Analytics and Revenue Management,\" in Proc. 2024 Int. Conf. Power, Energy, Control Transmiss. Syst. (ICPECTS), Chennai, India, 2024, pp. 1–6, doi: 10.1109/ICPECTS62210.2024.10780375.",
        "[9] R. Raut, A. B. Chandanshive, P. N. Gadkar, and E. Govardhan, \"Credit Card Fraud Detection Using Ensemble Modeling,\" in Proc. 2024 OPJU Int. Technol. Conf. (OTCON), Raigarh, India, 2024, pp. 1–6, doi: 10.1109/OTCON60325.2024.10687633.",
        "[10] I. D. Mienye and N. Jere, \"Deep Learning for Credit Card Fraud Detection: A Review of Algorithms, Challenges, and Solutions,\" IEEE Access, vol. 12, pp. 95081–95101, Jul. 2024, doi: 10.1109/ACCESS.2024.3426955.",
        "[11] K. Singhal, V. Singh, and A. Kaul, \"Smart Retail: Utilizing Machine Learning for Demand Prediction, Price Strategy, and Inventory Management,\" in Proc. 2024 IEEE 16th Int. Conf. Comput. Intell. Commun. Netw. (CICN), Malacca, Malaysia, 2024, pp. 1–6, doi: 10.1109/CICN63059.2024.10847534.",
        "[12] H. A. Chavan and P. P. Nitnaware, \"Smart Retail Solutions through Edge Computing and IoT Automation: Implementing Dynamic Pricing and Real-Time Customer Engagement,\" in Proc. 2025 IEEE 14th Int. Conf. Commun. Syst. Netw. Technol. (CSNT), Jabalpur, India, 2025, pp. 1–6, doi: 10.1109/CSNT64827.2025.10968920.",
        "[13] R. F. de Assis, W. de Paula Ferreira, A. F. Faria, L. A. Santa-Eulalia, M. Ouhimmou, and A. Gharbi, \"Optimising Warehouse Order Picking: Real Case Application in the Shoe Manufacturing Industry,\" IEEE Access, vol. 12, pp. 168434–168449, Nov. 2024, doi: 10.1109/ACCESS.2024.3497592.",
        "[14] E. Nugroho and G. Girsang, \"Three-Layer Multi-Objective VRP Solver: Modified AGE-MOEA-II, Greedy Split Delivery, and 2-opt,\" in Proc. 2025 5th Int. Conf. Electron. Elect. Eng. Intell. Syst. (ICE3IS), Yogyakarta, Indonesia, 2025, pp. 1–6, Electronic ISBN: 979-8-3315-8523-5. (IEEE Doc No: 10935612)",
        "[15] Y. Xiao, E. Xing, X. Sun, P. Wu, and D. Jiang, \"“Super Express-Courier” Plan: A Delivery Approach for Terminal Logistics-Stations Under Lean Management,\" in Proc. 2024 5th Int. Conf. Big Data, Artif. Intell. Internet Things Eng. (ICBAIE), Chongqing, China, 2024, pp. 1–6, IEEE Xplore: 10636402."
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(r)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    output_path = "docs/academic/FINAL_BLACK_BOOK.docx"
    doc.save(output_path)
    print(f"Successfully generated DOCX with real screenshots at: {output_path}")

if __name__ == "__main__":
    build_docx()
